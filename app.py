# -*- coding: utf-8 -*-
"""
################################################################################
#  O PREGADOR - SYSTEM OMEGA ENTERPRISE (V.50 - ULTIMATE BUILD)                #
#  STATUS: PRODUCTION | CRITICAL SYSTEM | REDUNDANT LOGIC                      #
#  --------------------------------------------------------------------------  #
#  ARQUITETURA DO SISTEMA:                                                     #
#  1. CORE BOOTLOADER: Verificação de integridade de sistema de arquivos.      #
#  2. SECURITY DAEMON: Criptografia AES-256 GCM (Grau Militar).                #
#  3. WORD PROCESSOR ENGINE: Módulo dedicado para manipulação de DOCX/PDF.     #
#  4. GENEVA DOCTRINE SCANNER: Algoritmo de análise semântica teológica.       #
#  5. SOUL ANALYTICS (PastoralMind): Análise comportamental e emocional.       #
#  6. NETWORK PROTOCOL: Camada de Rede Ministerial e Colaboração.              #
#  7. UX/UI RENDERER: Motor gráfico com animações CSS (Pulsing Cross).         #
################################################################################
"""
def inject_word_style():
    st.markdown("""
        <style>
            /* Expande o layout para quase 100% da tela */
            .main .block-container {
                max-width: 98% !important;
                padding-left: 1rem !important;
                padding-right: 1rem !important;
                padding-top: 1rem !important;
            }
            
            /* Estilização do Editor para parecer uma folha de papel */
            .ck-editor__editable {
                min-height: 700px !important;
                background-color: white !important;
                color: black !important;
                box-shadow: 0 0 10px rgba(0,0,0,0.1) !important;
                margin: 0 auto !important;
            }

            /* Ajuste para Mobile */
            @media (max-width: 768px) {
                .main .block-container {
                    padding-left: 0.5rem !important;
                    padding-right: 0.5rem !important;
                }
                .stButton button {
                    width: 100% !important;
                    margin-bottom: 5px;
                }
            }
        </style>
    """, unsafe_allow_html=True)

# Chame a função logo após o início
inject_word_style()
# ==============================================================================
# 00. IMPORTAÇÕES DE MÓDULOS DE SISTEMA
# ==============================================================================
import streamlit as st
import os
import sys
import time
import json
import base64
import math
import shutil
import random
import logging
import hashlib
import re
import sqlite3
import uuid
from datetime import datetime, timedelta
from io import BytesIO

# ===================================================================
# Page config & logging
# ===================================================================
st.set_page_config(
    page_title="O PREGADOR | SYSTEM OMEGA",
    layout="wide",
    page_icon="✝️",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://www.google.com',
        'Report a bug': "mailto:support@pregador.system",
        'About': "O PREGADOR V.50 Enterprise Edition"
    }
)

SYSTEM_ROOT = "Dados_Pregador_V31"
LOG_PATH = os.path.join(SYSTEM_ROOT, "System_Logs")
os.makedirs(LOG_PATH, exist_ok=True)

logging.basicConfig(
    filename=os.path.join(LOG_PATH, "system_audit_omega.log"),
    level=logging.INFO,
    format='[%(asctime)s] | [%(levelname)s] | MODULE: %(module)s | PROCESS: %(process)d | MSG: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logging.info(">>> INICIALIZAÇÃO DO SISTEMA OMEGA: CICLO STARTUP <<<")

# ===================================================================
# Feature detection for optional modules (editors, pdf, etc.)
# ===================================================================
GLOBAL_MODULES = {
    "CKEDITOR": False,
    "QUILL": False,
    "PLOTLY": False,
    "CRYPTO": False,
    "MAMMOTH": False,
    "HTML2DOCX": False,
    "REPORTLAB": False,
    "PYTHON-DOCX": False
}

try:
    from streamlit_ckeditor import st_ckeditor
    GLOBAL_MODULES["CKEDITOR"] = True
    logging.info("Dependência Carregada: Streamlit CKEditor")
except Exception as e:
    logging.warning(f"Dependência Falhou: CKEditor ({e})")

try:
    from streamlit_quill import st_quill
    GLOBAL_MODULES["QUILL"] = True
    logging.info("Dependência Carregada: Streamlit Quill")
except Exception as e:
    logging.warning(f"Dependência Falhou: Quill ({e})")

try:
    import plotly.graph_objects as go
    import plotly.express as px
    GLOBAL_MODULES["PLOTLY"] = True
    logging.info("Dependência Carregada: Plotly")
except Exception as e:
    logging.warning(f"Dependência Falhou: Plotly ({e})")

try:
    from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    GLOBAL_MODULES["CRYPTO"] = True
    logging.info("Dependência Carregada: Cryptography AES")
except Exception as e:
    logging.warning(f"Dependência Falhou: Cryptography ({e})")

try:
    import mammoth
    GLOBAL_MODULES["MAMMOTH"] = True
except Exception:
    pass

try:
    from html2docx import html2docx
    GLOBAL_MODULES["HTML2DOCX"] = True
except Exception:
    pass

try:
    from docx import Document
    GLOBAL_MODULES["PYTHON-DOCX"] = True
except Exception:
    pass

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    GLOBAL_MODULES["REPORTLAB"] = True
except Exception:
    pass

# ===================================================================
# Use modular core/visual/auth/utils/dashboard from app_modules
# ===================================================================
from app_modules.core import genesis_filesystem_integrity_check, DB_FILES, _read_json_safe, _write_json_atomic, DIRECTORY_STRUCTURE
from app_modules.visual import inject_visual_core
from app_modules.auth import AccessGate
from app_modules.utils import TextUtils
from app_modules import dashboard as dashboard_module

# Run genesis check and inject styles (single point)
genesis_filesystem_integrity_check()
inject_visual_core()

# ===================================================================
# Keep complex domain classes here if not moved to modules
# (WordProcessorEngine, GenevaDoctrineCore, PastoralAnalytics)
# ===================================================================
class WordProcessorEngine:
    def __init__(self, title, content_html, output_path):
        self.title = title
        self.content_html = content_html
        self.output_path = output_path
        self.clean_text = re.sub('<.*?>', '\n', content_html)

    def execute_docx_export(self):
        logging.info(f"Iniciando Exportação DOCX: {self.title}")
        # Mammoth
        if GLOBAL_MODULES.get("MAMMOTH"):
            try:
                import mammoth
                html_structure = f"<html><body><h1>{self.title}</h1>{self.content_html}</body></html>"
                result = mammoth.convert_to_docx(html_structure)
                with open(self.output_path, "wb") as f:
                    f.write(result.value)
                return True, "Documento gerado via Mammoth."
            except Exception as e:
                logging.error(f"Mammoth falhou: {e}")

        # html2docx
        if GLOBAL_MODULES.get("HTML2DOCX"):
            try:
                from html2docx import html2docx
                buf = html2docx(self.content_html, title=self.title)
                with open(self.output_path, "wb") as f:
                    f.write(buf.getvalue())
                return True, "Documento gerado via html2docx."
            except Exception as e:
                logging.error(f"html2docx falhou: {e}")

        # python-docx fallback
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            heading = doc.add_heading(self.title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for line in self.clean_text.split('\n'):
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)
            doc.add_page_break()
            doc.save(self.output_path)
            return True, "Documento gerado via python-docx."
        except Exception as e:
            try:
                txt_path = self.output_path.replace(".docx", ".txt")
                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(f"{self.title}\n\n{self.clean_text}")
                logging.error(f"Fallback txt gerado: {e}")
                return False, "Falha DOCX; salvo como TXT."
            except Exception as crit:
                return False, f"Falha crítica: {crit}"

    def execute_pdf_export(self):
        if not GLOBAL_MODULES.get("REPORTLAB"):
            return False, "ReportLab não instalado."
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            c = canvas.Canvas(self.output_path, pagesize=A4)
            width, height = A4
            c.setFillColorRGB(0.8, 0.7, 0.2)
            c.rect(0, height-50, width, 50, fill=1, stroke=0)
            c.setFillColorRGB(0,0,0)
            c.setFont("Helvetica-Bold", 18)
            c.drawCentredString(width/2, height-35, self.title)
            c.setFont("Helvetica", 11)
            y = height - 80
            for line in self.clean_text.split('\n'):
                for sub in [line[i:i+90] for i in range(0, len(line), 90)]:
                    if y < 50:
                        c.showPage()
                        y = height - 50
                        c.setFont("Helvetica", 11)
                    c.drawString(40, y, sub)
                    y -= 14
            c.save()
            return True, "PDF gerado."
        except Exception as e:
            logging.error(f"Erro PDF: {e}")
            return False, str(e)

class GenevaDoctrineCore:
    HERESY_DATABASE = {
        "prosperidade": "ALERTA: Viés de Prosperidade.",
        "decreto": "ALERTA: Antropocentrismo detectado.",
        "energia": "CUIDADO: Termo metafísico vago.",
        "universo": "ALERTA: Substituição panteísta.",
        "força": "NOTA: Verifique contexto.",
        "vibrar": "NOTA: Linguagem emocionalista detectada."
    }

    @staticmethod
    def scan(text):
        if not text: return []
        t = text.lower()
        found = []
        for k, msg in GenevaDoctrineCore.HERESY_DATABASE.items():
            if k in t: found.append(msg)
        return found

class PastoralAnalytics:
    @staticmethod
    def calculate_health_index():
        try:
            health_db_path = DB_FILES.get("HEALTH_DB") or os.path.join(SYSTEM_ROOT, "pastoral_health.json")
            db = _read_json_safe(health_db_path, {})
        except Exception:
            db = {}
        history = db.get("historico", []) if isinstance(db, dict) else []
        last = history[-10:] if history else []
        if not last:
            return 0.0, "Sem dados", "gray"
        scores = []
        for item in last:
            try:
                if isinstance(item, dict):
                    val = item.get("score", item.get("pontuacao", 0))
                else:
                    val = item
                scores.append(float(val))
            except Exception:
                scores.append(0.0)
        avg = round(max(0.0, min(100.0, (sum(scores)/len(scores)) if scores else 0.0)), 2)
        if avg >= 80: return avg, "Excelente", "green"
        if avg >= 60: return avg, "Bom", "lime"
        if avg >= 40: return avg, "Precisa atenção", "orange"
        return avg, "Crítico", "red"

# ===================================================================
# Session and Access Flow (uses AccessGate from app_modules.auth)
# ===================================================================
if "session_valid" not in st.session_state: st.session_state["session_valid"] = False
if "current_user" not in st.session_state: st.session_state["current_user"] = "GUEST"

if not st.session_state["session_valid"]:
    col_l, col_c, col_r = st.columns([1, 1.2, 1])
    with col_c:
        st.markdown("<br><br>", unsafe_allow_html=True)
        config = _read_json_safe(DB_FILES["CONFIG"], {})
        GOLD_HEX = config.get("theme_color", "#D4AF37")
        st.markdown(f"""
        <div class="login-container">
            <svg class="prime-logo" viewBox="0 0 100 100" xmlns="http://www.w3.org/2000/svg">
                <circle cx="50" cy="50" r="45" stroke="{GOLD_HEX}" stroke-width="2" fill="none" />
                <line x1="50" y1="20" x2="50" y2="80" stroke="{GOLD_HEX}" stroke-width="4" stroke-linecap="square" />
                <line x1="30" y1="40" x2="70" y2="40" stroke="{GOLD_HEX}" stroke-width="4" stroke-linecap="square" />
                <circle cx="50" cy="50" r="10" stroke="{GOLD_HEX}" stroke-width="1" fill="none" style="opacity:0.5"/>
            </svg>
            <div class="system-title">O PREGADOR</div>
            <div style="color: #666; font-size: 0.8rem; letter-spacing: 2px; margin-bottom: 20px;">
                SYSTEM OMEGA | ENTERPRISE EDITION V.50
            </div>
        </div>
        """, unsafe_allow_html=True)

        tab_access, tab_register = st.tabs(["🔒 ACESSO SEGURO", "📝 NOVO MEMBRO"])
        with tab_access:
            u = st.text_input("ID PASTORAL", placeholder="Digite seu identificador")
            p = st.text_input("CHAVE DE SEGURANÇA", type="password")
            if st.button("AUTENTICAR SISTEMA", use_container_width=True):
                if AccessGate.login_check(u, p):
                    st.session_state["session_valid"] = True
                    st.session_state["current_user"] = u.upper()
                    st.success("Credenciais validadas. Inicializando módulos...")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("ACESSO NEGADO: Credenciais não conferem.")

        with tab_register:
            st.info("Cadastro Local Habilitado")
            nu = st.text_input("Definir Novo ID", key="reg_u")
            np = st.text_input("Definir Senha", type="password", key="reg_p")
            if st.button("REGISTRAR CONTA"):
                success, msg = AccessGate.create_account(nu, np)
                if success: st.success(msg)
                else: st.error(msg)
    st.stop()

# ===================================================================
# Sidebar and navigation
# ===================================================================
with st.sidebar:
    st.markdown(f"""
    <div style="text-align:center; padding: 20px 0; border-bottom: 1px solid #222;">
        <h3 style="margin:0; font-size: 1.2rem;">Pastor</h3>
        <h1 style="margin:0; font-size: 1.8rem; line-height:1.2;">{st.session_state['current_user']}</h1>
    </div>
    """, unsafe_allow_html=True)

    health_score, health_status, health_color = PastoralAnalytics.calculate_health_index()
    st.markdown(f"""
    <div style="background:#111; padding:10px; border-radius:4px; margin: 15px 0; border:1px solid #222;">
        <small style="color:#888;">SAÚDE MINISTERIAL</small>
        <div style="display:flex; justify-content:space-between; align-items:center;">
            <span style="color:{health_color}; font-weight:bold;">{health_status}</span>
            <span style="color:{health_color}; font-size:1.2rem;">{health_score}%</span>
        </div>
        <div style="width:100%; height:4px; background:#222; margin-top:5px;">
            <div style="width:{health_score}%; height:100%; background:{health_color};"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    app_mode = st.radio(
        "Navegação do Sistema",
        ["Dashboard & Cuidado", "Gabinete de Preparação", "Rede Ministerial", "Biblioteca Digital", "Configurações"],
        label_visibility="collapsed"
    )

    st.markdown("---")
    if st.button("ENCERRAR SESSÃO", use_container_width=True):
        st.session_state["session_valid"] = False
        st.rerun()

# ===================================================================
# Main routing (Dashboard module is external)
# ===================================================================
if app_mode == "Dashboard & Cuidado":
    dashboard_module.render_dashboard()

elif app_mode == "Gabinete de Preparação":
    st.title("📝 Gabinete Pastoral Avançado")
    
    # Layout Proporcional: Coluna de arquivos menor (2), Editor maior (10)
    col_nav, col_work = st.columns([1.5, 8.5])
    
    with col_nav:
        st.markdown("### 🗂️ Acervo")
        all_files = [f for f in os.listdir(DIRECTORY_STRUCTURE["SERMONS"]) if f.endswith(".html")]
        all_files.sort(reverse=True)
        # Selectbox economiza mais espaço no mobile que Radio
        selected_file = st.selectbox("Escolher Documento", ["NOVO DOCUMENTO"] + all_files)
        
        st.divider()
        st.markdown("### 🛠️ Ações")
        # Botões de ação movidos para a lateral para liberar espaço vertical no editor
        c_save = st.container()
        c_word = st.container()
        c_pdf = st.container()
        c_scan = st.container()

    with col_work:
        active_content = ""
        active_title = ""
        if selected_file != "NOVO DOCUMENTO":
            file_path = os.path.join(DIRECTORY_STRUCTURE["SERMONS"], selected_file)
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    active_content = f.read()
                active_title = selected_file.replace(".html", "").replace("_", " ")
            except Exception as e:
                st.error(f"Erro de I/O: {e}")

        doc_title_input = st.text_input("TÍTULO DA MENSAGEM", value=active_title, placeholder="Ex: O Sermão do Monte")
        
        # --- ÁREA DO EDITOR ESTILO WORD ---
        final_editor_content = active_content

        if GLOBAL_MODULES.get("CKEDITOR"):
            # Configuração expandida do CKEditor
            custom_config = {
                'toolbar': [
                    ['heading', '|', 'bold', 'italic', 'link', 'bulletedList', 'numberedList', 'blockQuote', 'insertTable'],
                    ['undo', 'redo', '|', 'alignment', 'fontColor', 'fontSize', 'highlight']
                ],
                'height': 700, # Aumentado para parecer uma folha
                'width': '100%'
            }
            final_editor_content = st_ckeditor(
                value=active_content, 
                key="ck_editor_core", 
                config=custom_config
            )
        elif GLOBAL_MODULES.get("QUILL"):
            final_editor_content = st_quill(
                value=active_content, 
                html=True, 
                key="quill_editor_core",
                height=700
            )
        else:
            st.warning("Usando modo Texto Simples (Editores Visuais Indisponíveis)")
            final_editor_content = st.text_area("Editor", value=active_content, height=700)

        # Processamento de Dados (Lógica original preservada)
        clean_fname = TextUtils.sanitize_filename(doc_title_input if doc_title_input else "novo_sermao")

        with c_save:
            if st.button("💾 GRAVAR AGORA", use_container_width=True, type="primary"):
                save_path = os.path.join(DIRECTORY_STRUCTURE["SERMONS"], f"{clean_fname}.html")
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(final_editor_content)
                st.toast("Documento salvo!", icon="✅")
                time.sleep(1)
                st.rerun()

        with c_word:
            if st.button("📄 GERAR WORD (.docx)", use_container_width=True):
                out_path = os.path.join(DIRECTORY_STRUCTURE["SERMONS"], f"{clean_fname}.docx")
                processor = WordProcessorEngine(title=doc_title_input, content_html=final_editor_content, output_path=out_path)
                with st.spinner("Compilando..."):
                    status, msg = processor.execute_docx_export()
                if status:
                    st.success("Word Gerado!")
                    with open(out_path, "rb") as f:
                        st.download_button("BAIXAR DOCX", f, file_name=f"{clean_fname}.docx", use_container_width=True)

        with c_pdf:
            if st.button("📕 GERAR PDF", use_container_width=True):
                out_path = os.path.join(DIRECTORY_STRUCTURE["SERMONS"], f"{clean_fname}.pdf")
                processor = WordProcessorEngine(title=doc_title_input, content_html=final_editor_content, output_path=out_path)
                with st.spinner("Renderizando..."):
                    status, msg = processor.execute_pdf_export()
                if status:
                    st.success("PDF Pronto!")
                    with open(out_path, "rb") as f:
                        st.download_button("BAIXAR PDF", f, file_name=f"{clean_fname}.pdf", use_container_width=True)

        with c_scan:
            if st.button("🔍 SCAN DOUTRINÁRIO", use_container_width=True):
                alerts = GenevaDoctrineCore.scan(final_editor_content)
                if alerts:
                    for a in alerts: st.error(a)
                else:
                    st.success("Doutrina íntegra.")
        if GLOBAL_MODULES.get("CKEDITOR"):
            final_editor_content = st_ckeditor(value=active_content, key="ck_editor_core", height=600)
        elif GLOBAL_MODULES.get("QUILL"):
            final_editor_content = st_quill(value=active_content, html=True, key="quill_editor_core")
        else:
            st.warning("Editores Visuais Indisponíveis. Usando modo Raw Text.")
            final_editor_content = st.text_area("Editor Raw", value=active_content, height=600)

        st.markdown("---")
        st.markdown("### 🛠️ Processador de Engenharia Pastoral")
        c_save, c_word, c_pdf, c_scan = st.columns(4)
        clean_fname = TextUtils.sanitize_filename(doc_title_input if doc_title_input else "novo_sermao")

        if c_save.button("💾 GRAVAR (HTML)", use_container_width=True):
            save_path = os.path.join(DIRECTORY_STRUCTURE["SERMONS"], f"{clean_fname}.html")
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(final_editor_content)
            st.toast("Documento gravado com segurança no servidor.", icon="✅")
            time.sleep(1)
            st.rerun()

        if c_word.button("📄 EXPORTAR DOCX", use_container_width=True):
            out_path = os.path.join(DIRECTORY_STRUCTURE["SERMONS"], f"{clean_fname}.docx")
            processor = WordProcessorEngine(title=doc_title_input, content_html=final_editor_content, output_path=out_path)
            with st.spinner("Compilando Documento Word..."):
                status, msg = processor.execute_docx_export()
            if status:
                st.success(msg)
                with open(out_path, "rb") as f:
                    st.download_button("BAIXAR DOCX", f, file_name=f"{clean_fname}.docx")
            else:
                st.error(msg)

        if c_pdf.button("📕 EXPORTAR PDF", use_container_width=True):
            out_path = os.path.join(DIRECTORY_STRUCTURE["SERMONS"], f"{clean_fname}.pdf")
            processor = WordProcessorEngine(title=doc_title_input, content_html=final_editor_content, output_path=out_path)
            with st.spinner("Renderizando PDF Vectorial..."):
                status, msg = processor.execute_pdf_export()
            if status:
                st.success(msg)
                with open(out_path, "rb") as f:
                    st.download_button("BAIXAR PDF", f, file_name=f"{clean_fname}.pdf")
            else:
                st.error(msg)

        if c_scan.button("🔍 SCAN DOUTRINÁRIO", use_container_width=True):
            alerts = GenevaDoctrineCore.scan(final_editor_content)
            if alerts:
                with st.expander("⚠️ RELATÓRIO DE INCONSISTÊNCIA DETECTADO", expanded=True):
                    for a in alerts:
                        st.markdown(f"🔴 **{a}**")
            else:
                st.success("Análise completa: Nenhuma inconsistência detectada.")

elif app_mode == "Rede Ministerial":
    st.title("🤝 Rede de Conexão Pastoral")
    st.markdown("Central de inteligência compartilhada entre colaboradores.")
    feed_db = _read_json_safe(DB_FILES["NETWORK_FEED"], [])
    with st.expander("📡 TRANSMITIR NOVO CONTEÚDO", expanded=False):
        with st.form("feed_form"):
            p_title = st.text_input("Título do Recurso")
            p_author = st.text_input("Responsável", value=st.session_state["current_user"])
            p_link = st.text_input("Link Externo (YouTube/Drive/Vimeo)")
            p_body = st.text_area("Descrição Técnica / Teológica")
            if st.form_submit_button("PUBLICAR NA REDE"):
                if p_title and p_body:
                    packet = {"uuid": uuid.uuid4().hex, "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"), "title": p_title, "author": p_author, "link": p_link, "body": p_body}
                    feed_db.insert(0, packet)
                    _write_json_atomic(DB_FILES["NETWORK_FEED"], feed_db)
                    st.success("Pacote de dados transmitido.")
                    st.rerun()
                else:
                    st.warning("Dados incompletos.")

    st.divider()
    if not feed_db:
        st.info("Buffer de Rede Vazio. Aguardando transmissão.")
    for item in feed_db:
        with st.container():
            st.markdown(f"""
            <div class="tech-card">
                <div style="display:flex; justify-content:space-between;">
                    <h3 style="margin:0;">{item['title']}</h3>
                    <small>{item['timestamp']}</small>
                </div>
                <small style="color:var(--gold); font-weight:bold;">OPERADOR: {item['author']}</small>
                <hr style="border-color:#333;">
                <p>{item['body']}</p>
            </div>
            """, unsafe_allow_html=True)
            if "youtube" in (item.get('link') or ""):
                try: st.video(item['link'])
                except: st.warning("Erro de codec de vídeo.")
            elif item.get('link'):
                st.markdown(f"🔗 [Acessar Recurso Externo]({item['link']})")
            if st.button("REMOVER PACOTE", key=item['uuid']):
                feed_db.remove(item)
                _write_json_atomic(DB_FILES["NETWORK_FEED"], feed_db)
                st.rerun()

elif app_mode == "Biblioteca Digital":
    st.title("📚 Biblioteca & Acervo")
    col_up, col_idx = st.columns([1, 2])
    with col_up:
        st.markdown("### Ingestão de Ativos")
        f_up = st.file_uploader("Formatos: PDF, EPUB, DOCX", type=['pdf','epub','docx','txt'])
        if f_up:
            save_dest = os.path.join(DIRECTORY_STRUCTURE["LIBRARY_CACHE"], f_up.name)
            with open(save_dest, "wb") as f:
                f.write(f_up.getbuffer())
            st.success(f"Arquivo '{f_up.name}' armazenado no Cache.")
    with col_idx:
        st.markdown("### Índice Local")
        files = os.listdir(DIRECTORY_STRUCTURE["LIBRARY_CACHE"])
        if not files:
            st.warning("Cache vazio.")
        else:
            for file_name in files:
                f_path = os.path.join(DIRECTORY_STRUCTURE["LIBRARY_CACHE"], file_name)
                c1, c2 = st.columns([4, 1])
                c1.markdown(f"📑 **{file_name}**")
                with open(f_path, "rb") as f:
                    c2.download_button("BAIXAR", f, file_name=file_name, key=f"dl_{file_name}")

elif app_mode == "Configurações":
    st.title("⚙️ Painel de Controle do Sistema")
    tabs_conf = st.tabs(["🎨 Personalização UI", "🔧 Ferramentas de Manutenção", "ℹ️ Sobre o Sistema"])
    cfg = _read_json_safe(DB_FILES["CONFIG"], {})
    with tabs_conf[0]:
        st.subheader("Parâmetros Visuais")
        c_color, c_font = st.columns(2)
        new_color = c_color.color_picker("Cor Primária (Ouro)", cfg.get("theme_color"))
        new_font = c_font.selectbox("Tipografia", ["Inter", "Roboto", "Lato", "Cinzel"], index=0)
        if st.button("APLICAR MUDANÇAS DE TEMA"):
            cfg["theme_color"] = new_color
            cfg["font_family"] = new_font
            if _write_json_atomic(DB_FILES["CONFIG"], cfg):
                st.toast("Parâmetros regravados. Reinicialize a interface.")
                time.sleep(1)
                st.rerun()
    with tabs_conf[1]:
        st.subheader("Utilitários de Sistema")
        st.markdown("---")
        st.markdown("#### 1. Backup de Segurança")
        st.info("Gera um pacote .ZIP de todo o banco de dados.")
        if st.button("INICIAR PROTOCOLO DE BACKUP"):
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            fname = os.path.join(DIRECTORY_STRUCTURE["BACKUP_VAULT"], f"SystemBackup_{ts}")
            shutil.make_archive(fname, 'zip', SYSTEM_ROOT)
            st.success(f"Backup completo realizado em: {fname}.zip")
        st.markdown("---")
        st.markdown("#### 2. Auditoria de Logs")
        if os.path.exists(os.path.join(LOG_PATH, "system_audit_omega.log")):
            with open(os.path.join(LOG_PATH, "system_audit_omega.log"), "r") as log_f:
                lines = log_f.readlines()
            st.text_area("Dump de Logs de Auditoria", "".join(lines[-20:]), height=200)
            if st.button("LIMPAR LOGS"):
                 open(os.path.join(LOG_PATH, "system_audit_omega.log"), "w").close()
                 st.rerun()
    with tabs_conf[2]:
        st.markdown(
            """
            ### SYSTEM OMEGA V.50
            **Build:** Enterprise Gold Edition  
            **Status:** Stable Production  
            """
        )
        st.json(GLOBAL_MODULES)

# End of System
